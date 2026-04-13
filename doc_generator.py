import os
import re
from datetime import datetime
from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement


def replace_text_in_paragraphs(doc, replacements):
    for p in doc.paragraphs:
        for key, value in replacements.items():
            if key in p.text:
                p.text = p.text.replace(key, value)


def replace_text_in_tables(doc, replacements):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in replacements.items():
                        if key in paragraph.text:
                            paragraph.text = paragraph.text.replace(key, value)


def add_insert_paragraph_after():
    def insert_paragraph_after(self, text=None, style=None):
        new_p = OxmlElement("w:p")
        self._p.addnext(new_p)
        paragraph = Paragraph(new_p, self._parent)
        if text:
            paragraph.add_run(text)
        if style is not None:
            paragraph.style = style
        return paragraph

    Paragraph.insert_paragraph_after = insert_paragraph_after


def normalize_items(items):
    if items is None:
        return []

    if isinstance(items, str):
        text = items.strip()
        if not text:
            return []
        lines = [line.strip() for line in text.splitlines() if line.strip()]
    elif isinstance(items, list):
        lines = [str(item).strip() for item in items if str(item).strip()]
    else:
        lines = [str(items).strip()] if str(items).strip() else []

    cleaned = []
    for line in lines:
        line = re.sub(r'^[•\-\*]+\s*', '', line)
        line = re.sub(r'^[0-9٠-٩]+\s*[\.\)\-]\s*', '', line)
        cleaned.append(line.strip())

    return cleaned


def replace_placeholder_with_lines(doc, placeholder, items):
    clean_items = normalize_items(items)
    bullet_items = [f"• {item}" for item in clean_items if item]

    if not bullet_items:
        bullet_items = [""]

    for p in doc.paragraphs:
        if placeholder in p.text:
            p.text = p.text.replace(placeholder, bullet_items[0])

            current_p = p
            for item in bullet_items[1:]:
                new_p = current_p.insert_paragraph_after(item)
                current_p = new_p
            return

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if placeholder in p.text:
                        p.text = p.text.replace(placeholder, bullet_items[0])

                        current_p = p
                        for item in bullet_items[1:]:
                            new_p = current_p.insert_paragraph_after(item)
                            current_p = new_p
                        return


def add_heading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    p.alignment = 2  # right
    return p


def add_bullet_list(doc, items):
    for item in normalize_items(items):
        p = doc.add_paragraph(f"• {item}")
        p.alignment = 2


def add_plain_paragraph(doc, text):
    if text and str(text).strip():
        p = doc.add_paragraph(str(text).strip())
        p.alignment = 2


def add_responsibilities(doc, responsibilities):
    if not responsibilities:
        return

    for entry in responsibilities:
        party = entry.get("party", "").strip()
        details = normalize_items(entry.get("details", []))

        if party:
            p = doc.add_paragraph()
            r = p.add_run(f"{party}:")
            r.bold = True
            p.alignment = 2

        for d in details:
            p = doc.add_paragraph(f"• {d}")
            p.alignment = 2


def add_steps_table(doc, steps):
    clean_steps = steps if isinstance(steps, list) else []
    if not clean_steps:
        return

    add_heading(doc, "الإجراءات التفصيلية")

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"

    hdr = table.rows[0].cells
    hdr[0].text = "#"
    hdr[1].text = "المسؤولية"
    hdr[2].text = "الإجراء"
    hdr[3].text = "الأدوات و النماذج"

    for step in clean_steps:
        row = table.add_row().cells
        row[0].text = str(step.get("no", "")).strip()
        row[1].text = str(step.get("responsibility", "")).strip()
        row[2].text = str(step.get("procedure", "")).strip()
        row[3].text = str(step.get("tools", "")).strip()


def create_sop_doc_from_template(data):
    add_insert_paragraph_after()

    template_path = "templates/arabic_sop_template_placeholders.docx"
    doc = Document(template_path)

    os.makedirs("outputs", exist_ok=True)

    replacements = {
        "[عنوان الإجراء]": data.get("title", ""),
        "[القسم]": data.get("department", ""),
        "[الإصدار]": data.get("version", "01"),
        "[تاريخ الإصدار]": data.get("issue_date", ""),
        "[تاريخ التعميم]": data.get("circular_date", ""),
        "[جهة الاعتماد]": data.get("approval_entity", ""),
        "[تاريخ الاعتماد]": data.get("approval_date", ""),
        "[التوقيع]": data.get("signature", ""),
        "[المراجع الرئيسية]": data.get("main_references", ""),
        "[النماذج المستخدمة]": data.get("used_forms", ""),
        "[صاحب الإجراء]": data.get("procedure_owner", ""),
        "[إدارة الجودة]": data.get("quality_management", ""),
        "[المدير المفوض]": data.get("authorized_manager", ""),
        "[المقدمة]": data.get("introduction", ""),
        "[التعريف]": data.get("definition", ""),
    }

    replace_text_in_paragraphs(doc, replacements)
    replace_text_in_tables(doc, replacements)

    replace_placeholder_with_lines(doc, "[الأنظمة]", data.get("systems", []))
    replace_placeholder_with_lines(doc, "[الهدف]", data.get("objective", []))
    replace_placeholder_with_lines(doc, "[المطابقة]", data.get("matching_logic", []))
    replace_placeholder_with_lines(doc, "[نطاق التطبيق]", data.get("scope", []))
    replace_placeholder_with_lines(doc, "[شروط التصنيف]", data.get("classification_conditions", []))
    replace_placeholder_with_lines(doc, "[الضوابط الرقابية]", data.get("controls", []))
    replace_placeholder_with_lines(doc, "[المستندات المستخدمة]", data.get("documents_used", []))
    replace_placeholder_with_lines(doc, "[التقارير المطلوبة]", data.get("required_reports", []))

    steps = data.get("steps", [])

    # fill existing template rows if present
    for i in range(1, 21):
        step_data = steps[i - 1] if i <= len(steps) else {}
        step_replacements = {
            f"[رقم {i}]": str(step_data.get("no", "")),
            f"[المسؤولية {i}]": step_data.get("responsibility", ""),
            f"[الإجراء {i}]": step_data.get("procedure", ""),
            f"[الأدوات {i}]": step_data.get("tools", ""),
        }
        replace_text_in_paragraphs(doc, step_replacements)
        replace_text_in_tables(doc, step_replacements)

    # ALWAYS append full content at the end so nothing is lost
    doc.add_page_break()
    add_heading(doc, "النسخة التفصيلية الكاملة")

    add_heading(doc, "عنوان الإجراء")
    add_plain_paragraph(doc, data.get("title", ""))

    if data.get("introduction"):
        add_heading(doc, "المقدمة")
        add_plain_paragraph(doc, data.get("introduction", ""))

    if data.get("scope"):
        add_heading(doc, "نطاق التطبيق")
        add_bullet_list(doc, data.get("scope", []))

    if data.get("definition"):
        add_heading(doc, "التعريف")
        add_plain_paragraph(doc, data.get("definition", ""))

    if data.get("systems"):
        add_heading(doc, "الأنظمة المشاركة")
        add_bullet_list(doc, data.get("systems", []))

    if data.get("objective"):
        add_heading(doc, "هدف الإجراء")
        add_bullet_list(doc, data.get("objective", []))

    if data.get("matching_logic"):
        add_heading(doc, "منطق المطابقة")
        add_bullet_list(doc, data.get("matching_logic", []))

    if data.get("responsibilities"):
        add_heading(doc, "المسؤوليات")
        add_responsibilities(doc, data.get("responsibilities", []))

    if data.get("classification_conditions"):
        add_heading(doc, "شروط التصنيف")
        add_bullet_list(doc, data.get("classification_conditions", []))

    add_steps_table(doc, steps)

    if data.get("controls"):
        add_heading(doc, "الضوابط الرقابية")
        add_bullet_list(doc, data.get("controls", []))

    if data.get("documents_used"):
        add_heading(doc, "المستندات المستخدمة")
        add_bullet_list(doc, data.get("documents_used", []))

    if data.get("required_reports"):
        add_heading(doc, "التقارير المطلوبة")
        add_bullet_list(doc, data.get("required_reports", []))

    file_name = f"outputs/SOP_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(file_name)
    return file_name
